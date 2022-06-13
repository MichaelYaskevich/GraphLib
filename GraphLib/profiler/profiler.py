from pathlib import Path
from random import randint
from time import time
from docx import Document
from docx.shared import Pt
from docx.shared import Mm

from memory_profiler import memory_usage


from GraphLib.algorithms.pathSearch import \
    dijkstra, bellman_ford, algorithm_for_DAG, floyd
from GraphLib.generator.generator import generate_random_graph, generate_worst_case_graph_for_bellman_ford, \
    generate_best_case_graph_for_bellman_ford
from GraphLib.profiler.approximation import approximate
from GraphLib.profiler.statistic import Statistic
from GraphLib.profiler.visualization import visualize_time, visualize_memory

dummy_runs_count = 5
average_runs_count = 21


def profile_all_algorithms(path, min_size, max_size):
    algs_all_paths = {
        'dijkstra': dijkstra.find_shortest_paths,
        'bellman_ford': bellman_ford.find_shortest_paths,
        'algorithm_for_dag':
            algorithm_for_DAG.find_shortest_paths,
        'floyd': floyd.find_shortest_paths_from_source
    }
    time_data_dictionaries = []
    mem_data_dictionaries = []
    labels = []
    colors = ['black', 'red', 'green', 'blue']
    time_points_dicts = []
    time_conf_intervals = []
    mem_conf_intervals = []
    mem_points_dicts = []
    info_list = []
    args = generate_graphs(min_size, max_size)
    # args = generate_best_case_bf_graphs(min_size, max_size)
    # args = generate_worst_case_bf_graphs(min_size, max_size)
    for label, alg in algs_all_paths.items():
        labels.append(label)
        graph_sizes, time_statistics, memory_statistics, info = \
            profile(alg, args, label)
        info_list.append(info)

        mem_conf_intervals.append(
            {graph_sizes[i]: memory_statistics[i].confidence_interval
             for i in range(len(graph_sizes))})
        mem_points_dicts.append(
            {graph_sizes[i]: memory_statistics[i].avg
             for i in range(len(graph_sizes))})
        x, y = approximate(graph_sizes,
                           [memory_statistics[i].avg
                            for i in range(len(graph_sizes))])
        mem_data_dictionaries.append(
            {x[i]: y[i] for i in range(len(x))})

        time_conf_intervals.append(
            {graph_sizes[i]: time_statistics[i].confidence_interval
             for i in range(len(graph_sizes))})
        time_points_dicts.append(
            {graph_sizes[i]: time_statistics[i].avg
             for i in range(len(graph_sizes))})
        x, y = approximate(graph_sizes,
                           [time_statistics[i].avg
                            for i in range(len(graph_sizes))])
        time_data_dictionaries.append({x[i]: y[i] for i in range(len(x))})

    ROOT_DIR = Path(__file__).parent.parent.parent
    memory_path = Path(ROOT_DIR, 'GraphLib\\resources\\memory_image.png')
    time_path = Path(ROOT_DIR, 'GraphLib\\resources\\time_image.png')

    visualize_memory(mem_data_dictionaries, labels,
                     mem_points_dicts, mem_conf_intervals, colors, memory_path)
    visualize_time(time_data_dictionaries, labels,
                   time_points_dicts, time_conf_intervals, colors, time_path)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    doc.add_heading('Результаты исследования по времени и памяти')
    doc.add_paragraph(''.join(info_list))
    doc.add_picture(str(memory_path), width=Mm(170))
    doc.add_picture(str(time_path), width=Mm(170))
    doc.save(str(path))


def generate_graphs(min_size, max_size):
    result = []
    step = max(4, (max_size-min_size)//8)
    for i in range(max(min_size, 2), max_size, step):
        graph = generate_random_graph(i, 0, 1000)
        result.append((graph, randint(0, i - 1)))
    return result


def generate_worst_case_bf_graphs(min_size, max_size):
    result = []
    step = max(4, (max_size-min_size)//8)
    for i in range(max(min_size, 2), max_size, step):
        graph = generate_worst_case_graph_for_bellman_ford(i, 0, 1000)
        result.append((graph, randint(0, i - 1)))

    return result


def generate_best_case_bf_graphs(min_size, max_size):
    result = []
    step = max(4, (max_size-min_size)//8)
    for i in range(max(min_size, 2), max_size, step):
        graph = generate_best_case_graph_for_bellman_ford(i, 0, 1000)
        result.append((graph, i - 1))

    return result


def profile(func, args, label):
    time_statistics = []
    memory_statistics = []
    graph_sizes = []
    for graph, source in args:
        time_stat, memory_stat = get_profiling_results(func, graph, source)
        time_statistics.append(time_stat)
        memory_statistics.append(memory_stat)
        graph_sizes.append(len(graph.get_nodes()))
    info = make_report(label, time_statistics, memory_statistics)
    return graph_sizes, time_statistics, memory_statistics, info


def get_profiling_results(func, graph, source):
    for i in range(dummy_runs_count):
        func(graph, source)
    times, memory_us = get_times_and_memory_usage(func, graph, source)
    memory_statistic = Statistic(memory_us)
    time_statistic = Statistic(times)
    return time_statistic, memory_statistic


def get_times_and_memory_usage(func, graph, source):
    times = []
    for i in range(average_runs_count):
        current_run_start = time()
        list(func(graph, source))
        times.append(time() - current_run_start)
    memory_us = memory_usage(proc=lambda: list(func(graph, source)),
                             max_usage=False, backend="psutil",
                             include_children=True)
    return times, memory_us


def make_report(label, time_statistics, memory_statistics):
    time_statistic = combine_data(time_statistics)
    memory_statistic = combine_data(memory_statistics)
    info = f"""
    ------------------> RESULTS OF PROFILING {label} <------------------\n\n
    ------------------------- TIME INFO -------------------------
    Average time spent is {time_statistic.avg} seconds
    Minimum time spent is {time_statistic.minimum} seconds
    Maximum time spent is {time_statistic.maximum} seconds
    Standard deviation is {time_statistic.std_deviation} seconds
    Confidence interval (delta) for time """\
           f"""is {time_statistic.confidence_interval}'\n\n
    ------------------------ MEMORY INFO ------------------------
    Average memory usage is {memory_statistic.avg} MiB
    Minimum memory usage is {memory_statistic.minimum} MiB
    Maximum memory usage is {memory_statistic.maximum} MiB
    Standard deviation is {memory_statistic.std_deviation} MiB
    Confidence interval (delta) """\
           f"""for memory is {memory_statistic.confidence_interval}\n\n\n\n
    """

    print(info)

    return info


def combine_data(statistics):
    arr = []
    for stat in statistics:
        arr += stat.array
    return Statistic(arr)
