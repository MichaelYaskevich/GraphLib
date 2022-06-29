import matplotlib.pyplot as plt
import seaborn as sns

from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm

from main import log_func
from .visualization_data import VisualizationData


figsize = (16, 6)


def visualize_memory(data: VisualizationData, labels, colors, path):
    """
    Создает график результатов исследования по памяти.

    :param data: VisualizationData для памяти
    :param labels: надписи для кривых на графике
    :param colors: цвета кривых на графике
    :param path: путь, по которому надо сохранить график
    """

    plt.figure(figsize=figsize)
    plt.title("Memory performance research")
    plt.xlabel("data size, nodes count")
    plt.ylabel("memory, MiB")
    visualize(data, labels, colors, path)


def visualize_time(data: VisualizationData, labels, colors, path):
    """
    Создает график результатов исследования по времени.

    :param data: VisualizationData для времени
    :param labels: надписи для кривых на графике
    :param colors: цвета кривых на графике
    :param path: путь, по которому надо сохранить график
    """

    plt.figure(figsize=figsize)
    plt.title("Time performance research")
    plt.xlabel("data size, nodes count")
    plt.ylabel("time, seconds")
    visualize(data, labels, colors, path)


def visualize(data: VisualizationData, labels, colors, path):
    """
    Создает график результатов исследования.

    :param data: VisualizationData
    :param labels: надписи для кривых на графике
    :param colors: цвета кривых на графике
    :param path: путь, по которому надо сохранить график
    """

    for i in range(len(data.data_dictionaries)):
        sns.lineplot(data=data.data_dictionaries[i],
                     label=labels[i], color=colors[i])
        sns.scatterplot(data=data.points_dictionaries[i],
                        s=100, color=colors[i])
        for key, value in data.points_dictionaries[i].items():
            shift = data.confidence_intervals[i][key] / 2.0
            plt.plot([key, key], [value - shift, value + shift],
                     color=colors[i])
    plt.savefig(path)


def make_docx(info_list, labels, colors,
              mem_data: VisualizationData,
              time_data: VisualizationData) -> Document:
    """
    Создает docx документ результатов исследования по времени и памяти.

    :param info_list: подробные результаты профилирования
    :param labels: надписи для кривых на графике
    :param colors: цвета кривых на графике
    :param mem_data: VisualizationData для памяти
    :param time_data: VisualizationData для времени
    :return: docx документ результатов исследования по времени и памяти
    """

    root = Path(__file__).parent.parent.parent
    memory_path = Path(root, 'src', 'resources', 'memory_image.png')
    time_path = Path(root, 'src', 'resources', 'time_image.png')

    visualize_memory(mem_data, labels, colors, memory_path)
    visualize_time(time_data, labels, colors, time_path)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    doc.add_heading('Результаты исследования по времени и памяти')
    doc.add_paragraph(''.join(info_list))
    doc.add_picture(str(memory_path), width=Mm(170))
    doc.add_picture(str(time_path), width=Mm(170))

    return doc


def save_document(doc, path):
    """
    Сохраняет документ по указанному пути

    :param doc: docx документ
    :param path: абсолютный или относительный путь до папки или файла
    """

    if path.is_dir():
        doc.save(path / 'report.docx')
    else:
        if path.suffix != '.docx':
            doc.save(path.with_suffix('.docx'))
            log_func(f'Расширение файла было изменено '
                     f'с {path.suffix} на .docx')
        else:
            doc.save(path)
